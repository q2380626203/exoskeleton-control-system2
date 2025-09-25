#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
代码去注释并生成Word文档工具
用于外骨骼控制系统代码整理
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

def remove_comments(code, file_extension):
    """
    去掉代码中的注释
    支持C/C++, Arduino (.ino), Python等文件类型
    """
    if file_extension in ['.cpp', '.h', '.ino', '.c']:
        # C/C++风格注释
        # 去掉单行注释 //
        code = re.sub(r'//.*?$', '', code, flags=re.MULTILINE)
        # 去掉多行注释 /* */
        code = re.sub(r'/\*.*?\*/', '', code, flags=re.DOTALL)
    elif file_extension in ['.py']:
        # Python风格注释
        # 去掉单行注释 #
        code = re.sub(r'#.*?$', '', code, flags=re.MULTILINE)
        # 去掉多行字符串注释 """ """ 和 ''' '''
        code = re.sub(r'""".*?"""', '', code, flags=re.DOTALL)
        code = re.sub(r"'''.*?'''", '', code, flags=re.DOTALL)
    
    # 去掉空行和只含空白字符的行
    lines = code.split('\n')
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped:  # 只保留非空行
            clean_lines.append(line.rstrip())  # 去掉行尾空白
    
    return '\n'.join(clean_lines)

def count_lines(code):
    """统计代码行数"""
    return len([line for line in code.split('\n') if line.strip()])

def process_source_files(max_pages=60):
    """
    处理所有源代码文件
    返回: (总行数, 处理后的代码字典)
    max_pages: 最大页数限制
    """
    # 定义源文件列表（按优先级排序）
    source_files = [
        "main.c",
        "system_init.h", "system_init.c",
        "rs01_motor.h", "rs01_motor.c",
        "motor_web_control.h", "motor_web_control.c",
        "motor_web_html.h",
        "wifi_softap_module.h", "wifi_softap_module.c",
        "velocity_tracking_mode.h", "velocity_tracking_mode.c",
        "continuous_torque_velocity_mode.h", "continuous_torque_velocity_mode.c",
        "alternating_speed.h", "alternating_speed.c",
        "button_detector.h", "button_detector.c",
        "voice_module.h", "voice_module.c",
        "uart_motor_transmit.h", "uart_motor_transmit.c"
    ]
    
    processed_code = {}
    total_lines = 0
    
    # 源文件在main目录下
    main_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "main")
    if not os.path.exists(main_dir):
        main_dir = os.path.dirname(os.path.abspath(__file__))
    
    for filename in source_files:
        filepath = os.path.join(main_dir, filename)
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    original_code = f.read()
                
                # 获取文件扩展名
                _, ext = os.path.splitext(filename)
                
                # 去掉注释
                clean_code = remove_comments(original_code, ext)
                
                # 统计行数
                line_count = count_lines(clean_code)
                total_lines += line_count
                
                processed_code[filename] = {
                    'code': clean_code,
                    'lines': line_count,
                    'extension': ext
                }
                
                print(f"处理文件: {filename} - {line_count} 行")

                # 检查是否超过页数限制（包含文件标题）
                max_lines = max_pages * 50  # 每页50行
                estimated_total_with_titles = total_lines + len(processed_code)  # 加上已处理的文件标题
                if estimated_total_with_titles > max_lines:
                    print(f"已达到{max_pages}页限制({max_lines}行），剩余文件将被跳过")
                    break
                
            except Exception as e:
                print(f"处理文件 {filename} 时出错: {e}")
        else:
            print(f"文件不存在: {filename}")
    
    return total_lines, processed_code

def create_word_document(total_lines, processed_code, max_pages=60):
    """
    创建Word文档
    """
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    
    # 创建新的Word文档
    doc = Document()
    
    # 设置页面属性
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4纸高度
    section.page_width = Inches(8.27)    # A4纸宽度
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # 软件信息
    software_name = "外骨骼WiFi控制系统"
    version = "V1.0-ESP32-S3版"
    
    # 计算实际总行数（包括文件分隔符）
    lines_per_page = 50
    max_total_lines = max_pages * lines_per_page  # 最大行数限制
    
    # 先计算实际的总行数（包括文件标题）
    actual_total_lines = 0
    file_order = [
        "main.c",
        "system_init.h", "system_init.c",
        "rs01_motor.h", "rs01_motor.c",
        "motor_web_control.h", "motor_web_control.c",
        "motor_web_html.h",
        "wifi_softap_module.h", "wifi_softap_module.c",
        "velocity_tracking_mode.h", "velocity_tracking_mode.c",
        "continuous_torque_velocity_mode.h", "continuous_torque_velocity_mode.c",
        "alternating_speed.h", "alternating_speed.c",
        "button_detector.h", "button_detector.c",
        "voice_module.h", "voice_module.c",
        "uart_motor_transmit.h", "uart_motor_transmit.c"
    ]
    
    for filename in file_order:
        if filename in processed_code:
            # 每个文件有一个标题行
            actual_total_lines += 1
            # 加上文件的代码行数
            actual_total_lines += processed_code[filename]['lines']
    
    # 确保不超过最大页数
    actual_total_lines = min(actual_total_lines, max_total_lines)
    total_pages = min((actual_total_lines + lines_per_page - 1) // lines_per_page, max_pages)
    
    print(f"原始代码行数: {total_lines}")
    print(f"实际总行数（含文件标题）: {actual_total_lines}")
    print(f"实际页数: {total_pages}")
    
    # 设置页眉 - 创建两个段落实现左右对齐
    header = section.header
    
    # 清除默认页眉
    for para in header.paragraphs:
        para.clear()
    
    # 创建页眉表格来实现左右对齐
    from docx.oxml.shared import qn
    from docx.oxml import OxmlElement
    
    # 方法1：使用制表位实现左右对齐
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 设置制表位 - 更精确的右对齐位置
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Cm
    tab_stops = header_para.paragraph_format.tab_stops
    # 清除现有制表位
    tab_stops.clear_all()
    # 设置右对齐制表位在页面右边距前
    tab_stops.add_tab_stop(Cm(17.5), WD_TAB_ALIGNMENT.RIGHT)  # 更靠右的位置
    
    # 左侧软件名称
    left_run = header_para.add_run(f"{software_name} {version}")
    left_run.font.name = 'SimSun'
    left_run.font.size = Pt(10)
    
    # 制表符到右侧
    tab_run = header_para.add_run('\t第')
    tab_run.font.name = 'SimSun'
    tab_run.font.size = Pt(10)
    
    # 添加页码字段
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    # 页码运行
    page_run = header_para.add_run()
    page_run.font.name = 'SimSun'
    page_run.font.size = Pt(10)
    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)
    
    # 总页数 - 使用实际页数
    total_run = header_para.add_run(f'页，共{min(total_pages, max_pages)}页')
    total_run.font.name = 'SimSun'
    total_run.font.size = Pt(10)
    
    current_line_count = 0
    total_added_lines = 0  # 跟踪已添加的总行数

    # 按文件顺序处理代码（与上面保持一致）
    
    for filename in file_order:
        if filename not in processed_code:
            continue
            
        file_data = processed_code[filename]
        code_lines = file_data['code'].split('\n')

        # 检查是否超过最大页数限制
        if total_added_lines >= max_total_lines:
            print(f"已达到{max_pages}页限制，停止添加内容")
            break

        # 检查是否需要换页（只按50行规则，不为每个文件强制换页）
        if current_line_count >= lines_per_page:
            doc.add_page_break()
            current_line_count = 0
        
        # 添加文件名作为标题（不强制换页）
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(0)
        title_p.paragraph_format.line_spacing = 1.0
        title_run = title_p.add_run(f"// ========== {filename} ==========")
        title_run.font.name = 'Courier New'
        title_run.font.size = Pt(9)
        title_run.bold = True
        current_line_count += 1
        total_added_lines += 1

        # 检查页数限制
        if total_added_lines >= max_total_lines:
            print(f"已达到{max_pages}页限制，停止添加内容")
            break
        
        # 添加代码内容
        for line in code_lines:
            # 检查页数限制
            if total_added_lines >= max_total_lines:
                print(f"已达到{max_pages}页限制，停止添加内容")
                break

            # 只按50行规则换页
            if current_line_count >= lines_per_page:
                doc.add_page_break()
                current_line_count = 0

            # 添加代码行
            code_p = doc.add_paragraph()
            # 设置段落格式以减少行间距
            code_p.paragraph_format.space_before = Pt(0)
            code_p.paragraph_format.space_after = Pt(0)
            code_p.paragraph_format.line_spacing = 1.0  # 单倍行距

            code_run = code_p.add_run(line if line.strip() else ' ')  # 保持空行
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(8)  # 减小字体以容纳更多行

            current_line_count += 1
            total_added_lines += 1
    
    return doc

def main():
    """主函数"""
    print("开始处理外骨骼WiFi控制系统源代码...")

    # 处理源文件（限制60页）
    total_lines, processed_code = process_source_files(max_pages=60)
    
    if total_lines == 0:
        print("没有找到有效的源代码文件")
        return
    
    print(f"\n处理完成:")
    print(f"总文件数: {len(processed_code)}")
    print(f"总代码行数: {total_lines}")
    
    # 生成Word文档
    print("\n生成Word文档...")
    doc = create_word_document(total_lines, processed_code)
    
    # 保存文档
    import datetime
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f"外骨骼WiFi控制系统-源代码文档_60页版_{timestamp}.docx"
    doc.save(output_file)
    
    print(f"\nWord文档已生成: {output_file}")
    print("文档特点:")
    print("- 已去除所有注释")
    print("- 每页不少于50行代码")
    print("- 限制最多60页")
    print("- 页眉左上角：软件名称和版本号")
    print("- 页眉右上角：第X页，共Y页（动态页码）")
    print("- 使用Courier New等宽字体")
    print("- 制表符确保页码右对齐")
    print("\n提示: 如果页数显示仍不正确，请在Word中按F9刷新字段，或右键页码选择'更新域'。")

if __name__ == "__main__":
    main()